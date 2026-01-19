import aiohttp
import asyncio
import secrets
import sqlite3
import json
import os
from pathlib import Path

from aiogram import types
from aiogram.dispatcher import FSMContext
from aiogram.types import InlineKeyboardMarkup, InlineKeyboardButton, WebAppInfo

from config import BASE_DIR, CONTRACT_TOKENS_DB_PATH, logger_bot


path1 = str(Path(__file__).parents[2])


templates = {
    '1_1': "Авансовое соглашение от физического лица.docx",
    '1_2': "Авансовое соглашение от ИП.docx",
    '2': "Договор аренды.docx",
    '3_1': "Ипотека от физического лица.docx",
    '3_2': "Ипотека от ИП.docx",
    '3_3': "Ипотека от самозанятого.docx",
    '4_1': "Обмен от физического лица.docx",
    '4_2': "Обмен от ИП.docx",
    '4_3': "Обмен от самозанятого.docx",
    '5': "ПДКП_без поручительства.docx",
    '6_1': "Подбор от физического лица.docx",
    '6_2': "Подбор от ИП.docx",
    '6_3': "Подбор от самозанятого.docx",
    '7_1': "Продажа от физического лица.docx",
    '7_2': "Продажа от ИП.docx",
    '7_3': "Продажа от самозанятого.docx",
    '8': "Расторжение договора услуг.docx",
    '9_1': "Юридическое сопровождение от физического лица.docx",
    '9_2': "Юридическое сопровождение от ИП.docx",
    '9_3': "Юридическое сопровождение от самозанятого.docx",
    '10': "Соглашение о расторжении аванса.docx",
    '11': "Уведомление завышения-занижения краткое.docx",
    '12': "Уведомление завышения-занижения полное.docx",
    '13': "Уведомление о перепланировке(бланк).docx"
}

# Получаем внешний IP сервера


def get_external_ip() -> str:
    import requests
    try:
        return requests.get("https://api.ipify.org").text
    except Exception:
        return "127.0.0.1"  # fallback для отладки


# Отправка данных в FastAPI и генерация кнопки со ссылкой
async def send_passport_edit_link(message: types.Message, passport_data: dict, state: FSMContext):
    token = secrets.token_urlsafe(16)
    user_id = message.from_user.id
    payload = {
        "token": token,
        "user_id": user_id,
        **passport_data
    }
    ip_address = get_external_ip()
    fastapi_port = 80
    edit_url = f"https://neurochief.pro/edit/{token}"

    try:
        async with aiohttp.ClientSession() as session:
            async with session.post(
                f"https://neurochief.pro/api/save_passport_data1",
                json=payload
            ) as resp:
                if resp.status != 200:
                    text = await resp.text()
                    short_text = text[:800]
                    await message.answer(f"❌ Не удалось сохранить данные: {short_text}")
                    return

        # Создаем клавиатуру (старый способ для aiogram 2.25)
        keyboard = InlineKeyboardMarkup()
        keyboard.add(InlineKeyboardButton(
            "✏️ Изменить данные", web_app=WebAppInfo(url=edit_url)))
        await message.answer(
            "📝 Нажмите на кнопку ниже, чтобы отредактировать паспортные данные:",
            reply_markup=keyboard
        )

    except Exception as e:
        await message.answer(f"❌ Ошибка: {str(e)}")
        logger_bot.error(f"❌ Ошибка: {str(e)}")
    await wait_for_signal_and_run1(token, message, state)


async def send_passport_edit_link1(callback_query: types.CallbackQuery, passport_data: dict, state: FSMContext):
    token = secrets.token_urlsafe(16)
    user_id = callback_query.from_user.id
    payload = {
        "token": token,
        "user_id": user_id,
        **passport_data
    }
    ip_address = get_external_ip()
    fastapi_port = 80

    try:
        async with aiohttp.ClientSession() as session:
            async with session.post(
                # TODO Вернуть старую ссылку после тестов
                f"https://neurochief.pro/api/save_passport_data1",
                # f"http://localhost:8001/api/save_passport_data1",
                json=payload
            ) as resp:
                if resp.status != 200:
                    text = await resp.text()
                    await callback_query.message.answer(f"❌ Не удалось сохранить данные")
                    logger_bot.error(f"❌ Не удалось сохранить данные: {text}")
                    return

        # TODO Вернуть старую ссылку после тестов
        edit_url = f"https://neurochief.pro/edit/{token}"
        # edit_url = f"http://localhost:8001/edit/{token}"

        # Создаем клавиатуру
        keyboard = InlineKeyboardMarkup()
        keyboard.add(
            InlineKeyboardButton(
                "✏️ Изменить данные",
                web_app=WebAppInfo(url=edit_url)
            )
        )

        # TODO Вернуть после тестов
        # await callback_query.message.reply(f"Перейдите по ссылке: {edit_url}")
        await callback_query.message.answer(
            "📝 Нажмите на кнопку ниже, чтобы отредактировать паспортные данные:",
            reply_markup=keyboard
        )

        await wait_for_signal_and_run(token, callback_query, state)

    except Exception as e:
        logger_bot.error(f"❌ Ошибка: {e}")
        await callback_query.message.answer(f"❌ Ошибка: {str(e)}")


async def wait_for_signal_and_run1(token: str, message, state: FSMContext):
    user_id = message.from_user.id
    while True:
        await asyncio.sleep(3)  # каждые 3 секунды проверка
        logger_bot.info(f"Ожидаем заполнения данных для клиента {user_id}, {token}")
        try:
            with sqlite3.connect(CONTRACT_TOKENS_DB_PATH) as conn:
                cursor = conn.cursor()
                cursor.execute(
                    "SELECT Signal, data_json FROM tokens WHERE token = ?", (token,))
                result = cursor.fetchone()

                if result:
                    signal, data_json = result
                    if signal == 1:
                        # Парсим JSON
                        passport_data = json.loads(data_json)
                        await message.answer("✅ Данные обновлены! Генерирую договор...")
                        logger_bot.info("✅ Данные обновлены! Генерирую договор...")
                        # Запуск функции генерации договора
                        contract_path = await generate_contract(user_id, passport_data, state)
                        await message.answer_document(open(contract_path, 'rb'))
                        # Сбрасываем сигнал обратно в 0 (по желанию)
                        cursor.execute(
                            "UPDATE tokens SET Signal = 0 WHERE token = ?", (token,))
                        conn.commit()
                        break  # выходим из цикла после выполнения
        except Exception as e:
            logger_bot.error(f"Ошибка при проверке сигнала: {e}")
            break


async def wait_for_signal_and_run(token: str, callback_query, state: FSMContext):
    user_id = callback_query.from_user.id
    while True:
        await asyncio.sleep(10)  # каждые 3 секунды проверка
        try:
            with sqlite3.connect(CONTRACT_TOKENS_DB_PATH) as conn:
                cursor = conn.cursor()
                cursor.execute(
                    "SELECT Signal, data_json FROM tokens WHERE token = ?", (token,))
                result = cursor.fetchone()

                if result:
                    signal, data_json = result
                    if signal == 1:
                        # Парсим JSON
                        try:
                            passport_data = json.loads(data_json)
                            await callback_query.message.answer("✅ Данные обновлены! Генерирую договор...")
                            # Запуск функции генерации договора
                            contract_path = await generate_contract(user_id, passport_data, state)
                            await callback_query.message.answer_document(open(contract_path, 'rb'))
                            # Сбрасываем сигнал обратно в 0 (по желанию)
                            cursor.execute(
                                "UPDATE tokens SET Signal = 0 WHERE token = ?", (token,))
                            conn.commit()
                            break  # выходим из цикла после выполнения
                        except Exception as e:
                            await callback_query.message.answer("❌ Ошибка при обработке данных. Проверьте формат JSON.")
                            logger_bot.error(
                                f"❌ Ошибка при обработке данных. Проверьте формат JSON: {e}")
                            cursor.execute(
                                "UPDATE tokens SET Signal = 0 WHERE token = ?", (token,))
                            conn.commit()
                            break
        except Exception as e:
            logger_bot.error(f"Ошибка при проверке сигнала: {e}")
            break


async def generate_contract(user_id: int, passport_data: dict, state: FSMContext):
    from datetime import datetime
    from docx import Document

    # Разделяем риелтора и клиента (если нужно)
    rieltor_data, client_data = split_passport_data(passport_data)
    # Достаём doc_type
    data = await state.get_data()
    doc_type = data.get("doc_type")
    # словарь для соответствия business_type → суффикс
    business_suffix = {
        "Физическое лицо": "1",
        "ИП": "2",
        "Самозанятый": "3"
    }

    # проверяем, есть ли doc_type, который зависит от типа риелтора
    doc_types_with_business = {"1", "3", "4", "6", "7", "9"}

    if doc_type in doc_types_with_business:
        suffix = business_suffix.get(rieltor_data["business_type"])
        if suffix:
            doc_type = f"{doc_type}_{suffix}"

    # Путь к шаблону
    client_initials = f"{client_data['first_name'][0]}.{client_data['middle_name'][0]}." if client_data.get(
        'middle_name') else f"{client_data['first_name'][0]}."
    client_full_name = f"{client_data['last_name']} {client_initials}"
    template_path = os.path.join(
        BASE_DIR, "bot", "blanks", "templates", templates[doc_type])
    output_path = os.path.join(
        BASE_DIR, "bot", "contracts", f"{templates[doc_type].split('.')[0]}_{client_full_name}.docx"
    )

    rieltor_initials = f"{rieltor_data['first_name'][0]}.{rieltor_data['middle_name'][0]}." if rieltor_data.get(
        'middle_name') else f"{rieltor_data['first_name'][0]}."
    rieltor_full_name = f"{rieltor_data['last_name']} {rieltor_initials}"

    # Текущая дата
    current_date = datetime.now().strftime("%d.%m.%Y")

    # Формируем replacements динамически
    replacements = {f"{{{{{key}}}}}": str(value)
                    for key, value in passport_data.items()}

    # Добавляем вычисляемые поля
    replacements.update({
        "{{rieltor_name}}": rieltor_full_name,
        "{{client_name}}": client_full_name,
        "{{current_date}}": current_date,
    })

    # Открываем и заполняем шаблон
    doc = Document(template_path)

    # Функция для замены текста в параграфах и таблицах
    def replace_text_in_element(element):
        for key, value in replacements.items():
            if key in element.text:
                element.text = element.text.replace(key, value)

    # Заменяем в параграфах
    for paragraph in doc.paragraphs:
        replace_text_in_element(paragraph)

    # Заменяем в таблицах
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                replace_text_in_element(cell)

    # Сохраняем
    doc.save(output_path)
    return output_path


def split_passport_data(passport_data: dict) -> tuple[dict, dict]:
    rieltor_data = {}
    client_data = {}

    for key, value in passport_data.items():
        if key.startswith("rieltor_"):
            field = key.replace("rieltor_", "")
            rieltor_data[field] = value
        elif key.startswith("client_"):
            field = key.replace("client_", "")
            client_data[field] = value

    return rieltor_data, client_data
